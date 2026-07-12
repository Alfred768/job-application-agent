from io import BytesIO
from zipfile import ZipFile

from docx import Document

from job_agent.document_export import markdown_to_docx_bytes, tailor_docx_bytes


def test_markdown_to_docx_bytes_writes_basic_word_document():
    data = markdown_to_docx_bytes(
        "# Tailored Resume Draft\n\n## Base Resume\n\nGaoyi Wu\n\nBuilt FastAPI services."
    )

    with ZipFile(BytesIO(data)) as docx:
        names = set(docx.namelist())
        document_xml = docx.read("word/document.xml").decode("utf-8")

    assert "[Content_Types].xml" in names
    assert "word/document.xml" in names
    assert "Tailored Resume Draft" in document_xml
    assert "Base Resume" in document_xml
    assert "Gaoyi Wu" in document_xml
    assert "Built FastAPI services." in document_xml


def test_tailor_docx_preserves_source_and_reorders_existing_skills(tmp_path):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("GAOYI WU")
    document.add_paragraph("TECHNICAL SKILLS")
    document.add_paragraph("Backend: Redis, Python, FastAPI")
    document.add_paragraph("PROJECTS")
    document.save(source)

    tailored = tailor_docx_bytes(source, ["Python", "FastAPI"])
    output = tmp_path / "tailored.docx"
    output.write_bytes(tailored)

    source_doc = Document(source)
    tailored_doc = Document(output)
    assert source_doc.paragraphs[2].text == "Backend: Redis, Python, FastAPI"
    assert tailored_doc.paragraphs[2].text == "Backend: Python, FastAPI, Redis"
